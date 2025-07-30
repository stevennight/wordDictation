import os
import shutil
import json
from datetime import datetime
import docx

HISTORY_DIR = "history"
IMAGE_CACHE_DIR = "handwriting_cache"

def load_words_from_file(filepath):
    """Loads words from a .docx file."""
    try:
        doc = docx.Document(filepath)
        words = []
        if doc.tables:
            table = doc.tables[0]
            for row in table.rows:
                if len(row.cells) >= 2:
                    prompt1 = row.cells[0].text.strip()
                    answer1 = row.cells[1].text.strip()
                    if prompt1 and answer1:
                        words.append({"prompt": prompt1, "answer": answer1})
                if len(row.cells) >= 4:
                    prompt2 = row.cells[2].text.strip()
                    answer2 = row.cells[3].text.strip()
                    if prompt2 and answer2:
                        words.append({"prompt": prompt2, "answer": answer2})
        return words, f"成功加载 {len(words)} 个单词。"
    except Exception as e:
        return [], f"加载文件时出错: {e}"

def clear_handwriting_cache():
    """Clears the handwriting image cache directory."""
    if not os.path.exists(IMAGE_CACHE_DIR):
        os.makedirs(IMAGE_CACHE_DIR)
    else:
        for f in os.listdir(IMAGE_CACHE_DIR):
            os.remove(os.path.join(IMAGE_CACHE_DIR, f))

import uuid

HISTORY_INDEX_FILE = os.path.join(HISTORY_DIR, "history_index.json")

def load_history_index():
    if not os.path.exists(HISTORY_INDEX_FILE):
        return []
    with open(HISTORY_INDEX_FILE, 'r', encoding='utf-8') as f:
        return json.load(f)

def save_history_index(index):
    with open(HISTORY_INDEX_FILE, 'w', encoding='utf-8') as f:
        json.dump(index, f, ensure_ascii=False, indent=4)

def save_history(results, stats, word_file_path, config, is_retry=False):
    """Saves the dictation session to the history."""
    if not os.path.exists(HISTORY_DIR):
        os.makedirs(HISTORY_DIR)

    current_timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    history_timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    random_str = uuid.uuid4().hex[:8]
    history_filename_base = f"{history_timestamp}_{random_str}"
    history_json_path = os.path.join(HISTORY_DIR, f"{history_filename_base}.json")
    history_image_dir = os.path.join(HISTORY_DIR, history_filename_base)

    if not os.path.exists(history_image_dir):
        os.makedirs(history_image_dir)

    history_data = {
        "results": []
    }

    for i, res in enumerate(results):
        original_image_path = res.get("original_image_path")
        annotated_image_path = res.get("annotated_image_path")

        if original_image_path and os.path.exists(original_image_path):
            shutil.copy(original_image_path, history_image_dir)
        
        if annotated_image_path and os.path.exists(annotated_image_path):
            shutil.copy(annotated_image_path, history_image_dir)

        history_data["results"].append({
            "prompt": res["prompt"],
            "answer": res["answer"],
            "correct": res["correct"],
            "original_image_path": os.path.join(history_image_dir, os.path.basename(original_image_path)) if original_image_path else None,
            "annotated_image_path": os.path.join(history_image_dir, os.path.basename(annotated_image_path)) if annotated_image_path else None
        })

    with open(history_json_path, 'w', encoding='utf-8') as f:
        json.dump(history_data, f, ensure_ascii=False, indent=4)

    index = load_history_index()
    if is_retry and ' (' in word_file_path and ')' in word_file_path:
        # 从 "文件名 (时间)" 格式中提取文件名和时间
        word_file_name = word_file_path[:word_file_path.rfind(' (')]
        original_timestamp = word_file_path[word_file_path.rfind('(') + 1:word_file_path.rfind(')')]
    else:
        word_file_name = os.path.basename(word_file_path) if word_file_path else "重做错题"
        original_timestamp = current_timestamp

    index.append({
        "filename": f"{history_filename_base}.json",
        "word_file_name": word_file_name,
        "timestamp": current_timestamp,
        "original_timestamp": original_timestamp if is_retry else current_timestamp,
        "is_retry": is_retry,
        "stats": stats
    })

    max_history = config.get("max_history_size", 10)
    print(f"当前历史记录数量: {len(index)}, 最大允许数量: {max_history}")
    enforce_history_limit(index, max_history)

def delete_history_record(filename):
    """Deletes a specific history record's files and its index entry."""
    index = load_history_index()
    
    # Find and remove the record from the index
    updated_index = [record for record in index if record['filename'] != filename]
    
    if len(updated_index) < len(index):
        save_history_index(updated_index)
        
        # Delete the associated files
        json_path = os.path.join(HISTORY_DIR, filename)
        if os.path.exists(json_path):
            os.remove(json_path)

        image_dir = os.path.join(HISTORY_DIR, os.path.splitext(filename)[0])
        if os.path.exists(image_dir):
            shutil.rmtree(image_dir)
        
        return True
    return False

def enforce_history_limit(index, max_size):
    """Ensures the number of history records does not exceed the max size."""
    if len(index) <= max_size:
        save_history_index(index)
        return

    # Sort by timestamp to find the oldest records
    index.sort(key=lambda x: x.get('timestamp', ''))

    # Determine which records to delete and which to keep
    num_to_delete = len(index) - max_size
    records_to_delete = index[:num_to_delete]
    remaining_records = index[num_to_delete:]

    # Delete the files of the oldest records
    for record in records_to_delete:
        delete_history_record(record['filename'])

    # Save the updated index
    save_history_index(remaining_records)